import os
import json
import uuid
import re
import logging
import docx
import pandas as pd
from PIL import Image
import pytesseract
from datetime import datetime
from pypdf import PdfReader
from django.shortcuts import render
from django.http import JsonResponse
from django.views.decorators.csrf import csrf_exempt
from django.conf import settings
from google.generativeai import GenerativeModel
import google.generativeai as genai
from .models import FileUpload,Chat,WebPage
from .tasks import run_spider_task
from urllib.parse import urlparse


from ahanaapp.models import WebPage



from pdf2image import convert_from_path  # To convert PDF pages to images
import io


GENAI_API_KEY = "AIzaSyD-_BehWwc-B_5BexpncXPlUj83_eGxf6o"  # keep this secret
genai.configure(api_key=GENAI_API_KEY)
model = GenerativeModel("gemini-2.0-flash")

import openpyxl

# ✅ Constants
UPLOAD_FOLDER = os.path.join(settings.MEDIA_ROOT, "uploads")
poppler_path = r"C:\Users\user\Downloads\Release-24.08.0-0\poppler-24.08.0\Library\bin"
EXCEL_FILE = os.path.join(settings.BASE_DIR, "chat-history.xlsx")


# ✅ Ensure upload folder exists
os.makedirs(UPLOAD_FOLDER, exist_ok=True)

# ✅ Configure logging
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')

# ✅ Gemini API configuration


# ✅ Load static responses and URLs
config_path = os.path.join(settings.BASE_DIR, "config.json")
try:
    with open(config_path, "r", encoding="utf-8") as f:
        config = json.load(f)
    company_website = config.get("company_websites", {})
    responses = config.get("responses", {})
except Exception as e:
    logging.error(f"Failed to load config.json: {e}")
    company_website, responses = {}, {}

def get_file_text_data(request):
    if "file_text_data" not in request.session:
        request.session["file_text_data"] = {}
    return request.session["file_text_data"]


# ✅ File type validation
def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in {"pdf", "docx", "txt", "jpg", "jpeg", "png", "csv", "xlsx"}

# ✅ Extract text from various file types
def extract_text_from_file(file_path, file_ext):
    try:
        extracted_text = ""

        file_ext = file_ext.lower()  # Always lowercase

        if file_ext == "pdf":
            # Step 1: Extract normal text from PDF
            pdf_reader = PdfReader(file_path)
            for page in pdf_reader.pages:
                text = page.extract_text()
                if text:
                    extracted_text += text.strip() + "\n"

            if extracted_text.strip():  # If normal text is found, return
                logging.info(f"Extracted normal text from PDF: {file_path}")
                return extracted_text.strip()

            # Step 2: If no text found, do OCR
            logging.info(f"No normal text found, doing OCR on PDF images: {file_path}")
            images = convert_from_path(file_path, poppler_path=poppler_path)
            ocr_text = ""
            for img in images:
                ocr_text += pytesseract.image_to_string(img).strip() + "\n"

            return ocr_text.strip() if ocr_text.strip() else "No readable text found in the document."

        elif file_ext == "docx":
            doc = docx.Document(file_path)
            extracted_text = [para.text.strip() for para in doc.paragraphs if para.text.strip()
            ] + [
                " | ".join([cell.text.strip() for cell in row.cells if cell.text.strip()])
                for table in doc.tables for row in table.rows
            ] + [
                f"HEADER: {para.text.strip()}" for section in doc.sections if section.header 
                for para in section.header.paragraphs if para.text.strip()
            ] + [
                f"FOOTER: {para.text.strip()}" for section in doc.sections if section.footer
                for para in section.footer.paragraphs if para.text.strip()
            ]

            return "\n".join(extracted_text) if extracted_text else "No text found in the document."


        elif file_ext == "txt":
            with open(file_path, "r", encoding="utf-8") as f:
                return f.read().strip()

        elif file_ext == "xlsx":
            df = pd.read_excel(file_path)
            return df.to_string(index=False).strip()

        elif file_ext == "csv":
            df = pd.read_csv(file_path)
            return df.to_string(index=False).strip()

        elif file_ext in {"jpg", "jpeg", "png"}:
            image = Image.open(file_path)
            ocr_text = pytesseract.image_to_string(image)
            return ocr_text.strip() if ocr_text.strip() else "No readable text found in the image."

        return "Unsupported file type."

    except Exception as e:
        logging.error(f"Error extracting text from {file_path}: {e}")
        return f"Error extracting text: {e}"

# ✅ Save chat to database
def save_chat_to_db(session_id, user_message, bot_response):
    chat_entry = Chat(
        session_id=session_id,
        user_message=user_message,
        bot_response=bot_response
    )
    try:
        chat_entry.save()
        logging.info("Chat saved to database.")
    except Exception as e:
        logging.error(f"Error saving chat to database: {e}")


# ✅ Generate or get session ID
def get_session_id(request):
    if "session_id" not in request.session:
        request.session["session_id"] = str(uuid.uuid4())
    return request.session["session_id"]

# ✅ Home View
def home(request):
    request.session.pop("session_id", None)
    request.session.pop("file_text_data", None)
    return render(request, "index.html")

# ✅ File Upload Handler


# ✅ File Upload Handler
@csrf_exempt
def upload_file(request):
    if request.method == "POST":
        uploaded_file = request.FILES.get("file")
        if not uploaded_file or not allowed_file(uploaded_file.name):
            return JsonResponse({"error": "Invalid file type."}, status=400)

        file_ext = uploaded_file.name.split('.')[-1].lower()
        file_path = os.path.join(UPLOAD_FOLDER, uploaded_file.name)

        try:
            with open(file_path, "wb") as f:
                for chunk in uploaded_file.chunks():
                    f.write(chunk)
        except Exception as e:
            logging.error(f"Failed to save file: {e}")
            return JsonResponse({"error": f"Failed to save file: {e}"}, status=500)

        text = extract_text_from_file(file_path, file_ext)
        file_data = get_file_text_data(request)
        file_data[os.path.basename(uploaded_file.name)] = text
        request.session.modified = True  # Mark session as modified so it saves

        try:
            chat = model.start_chat(history=[{
                "role": "user",
                "parts": [f"Summarize the document:\n\n{text[:9000]}"]
            }])
            summary = chat.send_message("Summarize the document.").text or "Gemini couldn't summarize."
        except Exception as e:
            logging.error(f"Gemini error: {e}")
            summary = f"Gemini API Error: {e}"

        session_id = get_session_id(request)
        response = f"File '{uploaded_file.name}' uploaded successfully!\n\nSummary:\n{summary}"

        # Save file upload data to database
        file_upload_entry = FileUpload(
            session_id=session_id,
            file_name=uploaded_file.name,
            extracted_text=text,
            bot_response=response
        )
        try:
            file_upload_entry.save()
            logging.info("File upload saved to database.")
        except Exception as e:
            logging.error(f"Error saving file upload to database: {e}")

        save_chat_to_db(session_id, f"Uploaded file: {uploaded_file.name}", response)

        return JsonResponse({"response": response, "file_name": uploaded_file.name, "extracted_text": text})

@csrf_exempt
def chatbot(request):
    if request.method == "POST":
        try:
            data = json.loads(request.body)
            user_msg = data.get("message", "").strip()
            history = data.get("history", [])
        except json.JSONDecodeError:
            return JsonResponse({"error": "Invalid JSON."}, status=400)

        if not user_msg:
            default = "Hi! I am Ahana, your chatbot. How can I assist you today?"
            save_chat_to_db(get_session_id(request), user_msg, default)
            return JsonResponse({"response": default})

        file_data = get_file_text_data(request)

        # Check if user asked to compare two files
        compare_files = []
        for fname in file_data:
            if fname.lower() in user_msg.lower():
                compare_files.append(fname)
        if "compare" in user_msg.lower() and len(compare_files) >= 2:
            fname1, fname2 = compare_files[0], compare_files[1]
            ftext1, ftext2 = file_data[fname1], file_data[fname2]

            # Create a more refined comparison prompt
            comparison_prompt = f"""
            Compare the following two files based on the points below:

            **1. File: {fname1}**  
            - **Content:**  
                - {ftext1[:15000]}
            - **Objective:**  
                - Understanding the purpose behind {fname1}
            - **Subject Matter:**  
                - Main topics covered in {fname1}

            **2. File: {fname2}**  
            - **Content:**  
                - {ftext2[:15000]}
            - **Objective:**  
                - Understanding the purpose behind {fname2}
            - **Subject Matter:**  
                - Main topics covered in {fname2}

            ---

            **Similarities**
            - **File Type:**  
                - Both files are [file_type] (e.g., PDF, Word, etc.)
            - **Content Focus:**  
                - Both discuss [common_topic]
            - **Target Audience:**  
                - Both files are aimed at [intended_audience]
            - **Overall Summary:**  
                - The files have [similarity_description], but differ in details.

            ---

            **Differences**
            - **File Type:**  
                - {fname1} is [type], whereas {fname2} is [type].
            - **Content Focus:**  
                - {fname1} focuses on [topic], while {fname2} focuses on [topic].
            - **Target Audience:**  
                - {fname1} is intended for [audience], while {fname2} is aimed at [audience].
            - **Overall Summary:**  
                - {fname1} is [description], whereas {fname2} is [description].

            **Extracted File Contents**
            **1. File: {fname1}**  
            - {ftext1[:15000]}

            **2. File: {fname2}**  
            - {ftext2[:15000]}
            """ 

            # Sending the comparison prompt to the model
            try:
                formatted_history = [{"role": h["role"], "parts": [h["content"]]} for h in history]
                formatted_history.append({"role": "user", "parts": [comparison_prompt]})

                chat = model.start_chat(history=formatted_history)
                gemini_reply = chat.send_message(comparison_prompt)
                response_text = gemini_reply.text or "Sorry, I couldn't generate the comparison."
            except Exception as e:
                logging.error(f"Gemini API error during comparison: {e}")
                response_text = f"Gemini API error during comparison: {e}"

            save_chat_to_db(get_session_id(request), user_msg, response_text)
            return JsonResponse({"response": response_text})

        # Handle extracting text if user just asks about a single file
        for fname, ftext in file_data.items():
            if fname.lower() in user_msg.lower():
                response = f"Extracted text from {fname}:\n\n{ftext[:1000]}..."
                save_chat_to_db(get_session_id(request), user_msg, response)
                return JsonResponse({"response": response})

        # Handle responses from predefined keywords
        for keyword, url in company_website.items():
            if re.search(rf"\b{keyword}\b", user_msg, re.IGNORECASE):
                response = f"Redirecting to {keyword} website..."
                if keyword == "support":
                    response += " Contact support@ahanait.com for help."
                save_chat_to_db(get_session_id(request), user_msg, response)
                return JsonResponse({"response": response, "url": url})

        for keyword, reply in responses.items():
            if re.search(rf"\b{keyword}\b", user_msg, re.IGNORECASE):
                save_chat_to_db(get_session_id(request), user_msg, reply)
                return JsonResponse({"response": reply})

        today = datetime.now()
        user_msg_lower = user_msg.lower()

        if "today" in user_msg_lower and "date" in user_msg_lower:
            response_text = today.strftime('%d %B %Y')
            save_chat_to_db(get_session_id(request), user_msg, response_text)
            return JsonResponse({"response": response_text})

        if "yesterday" in user_msg_lower and "date" in user_msg_lower:
            yesterday = today - pd.Timedelta(days=1)
            response_text = yesterday.strftime('%d %B %Y')
            save_chat_to_db(get_session_id(request), user_msg, response_text)
            return JsonResponse({"response": response_text})

        if "tomorrow" in user_msg_lower and "date" in user_msg_lower:
            tomorrow = today + pd.Timedelta(days=1)
            response_text = tomorrow.strftime('%d %B %Y')
            save_chat_to_db(get_session_id(request), user_msg, response_text)
            return JsonResponse({"response": response_text})
            
            
        session_id = get_session_id(request)
        url_match = re.search(r'https?://[^\s]+', user_msg)
        low_msg = user_msg.lower()

        # 🌐 Trigger crawler only if a new site is detected
        if url_match:
            parsed = urlparse(url_match.group())
            website_url = f"{parsed.scheme}://{parsed.netloc}"
            existing_pages = WebPage.objects.filter(url__startswith=website_url)

            if not existing_pages.exists():
                run_spider_task.delay(website_url)
                request.session["last_crawled_url"] = website_url
                response = f"🔍 Crawling started for {website_url}. Please wait a minute before asking questions."
                save_chat_to_db(session_id, user_msg, response)
                return JsonResponse({"response": response})
            else:
                request.session["last_crawled_url"] = website_url

        # ✅ Only trigger website logic if the message contains website-related keywords
        last_crawled_url = request.session.get("last_crawled_url", "")
        website_keywords = ["website", "site", "page", "crawl", "link", "summary", "internal"]
        is_website_question = any(kw in low_msg for kw in website_keywords)

        if last_crawled_url and is_website_question:
            pages = WebPage.objects.filter(url__startswith=last_crawled_url)
            if not pages.exists():
                response = "⏳ Crawling in progress. Please wait and ask again shortly."
                save_chat_to_db(session_id, user_msg, response)
                return JsonResponse({"response": response})

            # 🔹 Summary
            if "summary" in low_msg or "what is this website" in low_msg:
                full_text = " ".join(page.content for page in pages if page.content)
                prompt = f"Please summarize the following website content:\n\n{full_text[:30000]}"
                try:
                    chat = model.start_chat(history=[])
                    gemini_reply = chat.send_message(prompt)
                    summary_text = gemini_reply.text or "Sorry, I couldn't generate the summary."
                    save_chat_to_db(session_id, user_msg, summary_text)
                    return JsonResponse({"response": summary_text})
                except Exception as e:
                    logging.error(f"Gemini API error during summarization: {e}")
                    return JsonResponse({"response": "Error during summary generation."})

            # 🔹 Link count
            elif "how many links" in low_msg:
                internal_link_count = sum(len(page.internal_links or []) for page in pages)
                external_link_count = sum(len(page.external_links or []) for page in pages)
                response_text = (
                f"🔗 The website contains:\n"
                f"📌 {internal_link_count} **internal links** (within the same site)\n"
                f"🌍 {external_link_count} **external links** (leading to other websites)"
                )
                save_chat_to_db(session_id, user_msg, response_text)
                return JsonResponse({"response": response_text})

            # 🔹 Show links
            elif ("provide" in low_msg and "link" in low_msg) or "show me links" in low_msg:
                all_internal_links = []
                all_external_links = []
                for page in pages:
                    if page.internal_links:
                        all_internal_links.extend(page.internal_links)
                    if page.external_links:
                        all_external_links.extend(page.external_links)

                all_internal_links = list(set(all_internal_links))
                all_external_links = list(set(all_external_links))

                if not all_internal_links and not all_external_links:
                    response_text = "⚠️ No links found yet. Crawling may still be in progress."
                else:
                    response_text = "🔗 Here are some links found on the website:\n\n"
                    if all_internal_links:
                        internal_links_to_show = all_internal_links[:30]
                        response_text += "**📌 Internal Links:**\n" + "\n".join(internal_links_to_show) + "\n\n"
                    if all_external_links:
                        external_links_to_show = all_external_links[:30]
                        response_text += "**🌍 External Links:**\n" + "\n".join(external_links_to_show)

                save_chat_to_db(session_id, user_msg, response_text)
                return JsonResponse({"response": response_text})

                # 🔹 General QA from website
            else:
                question_keywords = set(re.findall(r'\w+', low_msg))
                ranked_pages = []
                for page in pages:
                    content_lower = page.content.lower()
                    score = sum(content_lower.count(word) for word in question_keywords)
                    if score > 0:
                        ranked_pages.append((score, page))

                if not ranked_pages:
                    response = "I couldn't find relevant information in the website content."
                    save_chat_to_db(session_id, user_msg, response)
                    return JsonResponse({"response": response})

                ranked_pages.sort(reverse=True, key=lambda x: x[0])
                top_pages = ranked_pages[:3]
                combined_content = ""
                token_limit = 30000

                for _, page in top_pages:
                    text_block = f"\n\nPage: {page.url}\n{page.content}"
                    if len(combined_content) + len(text_block) < token_limit:
                        ombined_content += text_block
                    else:
                        break

                prompt = f"""Answer the following question based only on the content below from the website pages:
                User Question: {user_msg}
                Relevant Website Content: {combined_content}"""
                try:
                    chat = model.start_chat(history=[])
                    gemini_reply = chat.send_message(prompt)
                    answer = gemini_reply.text or "Sorry, I couldn't find a clear answer in the relevant content."
                    save_chat_to_db(session_id, user_msg, answer)
                    return JsonResponse({"response": answer})
                except Exception as e:
                    logging.error(f"Gemini API error during QA: {e}")
                    return JsonResponse({"response": "Error during question answering."})

        # Default: normal conversation using all files combined
        matched_fname = None
        for fname in file_data:
            if fname.lower() in user_msg.lower():
                matched_fname = fname
                break

        if matched_fname:
            ftext = file_data[matched_fname]
            file_context = f"\n\nHere is the content of the file '{matched_fname}':\n{ftext[:40000]}"
        else:
            combined_text = "\n\n".join([f"File: {fname}\n{ftext[:10000]}" for fname, ftext in file_data.items()])
            file_context = f"\n\nCombined file content:\n{combined_text[:40000]}"

        final_prompt = f"{file_context}\n\nUser question: {user_msg}" if file_context else user_msg

        try:
            formatted_history = [{"role": h["role"], "parts": [h["content"]]} for h in history]
            formatted_history.append({"role": "user", "parts": [final_prompt]})

            chat = model.start_chat(history=formatted_history)
            gemini_reply = chat.send_message(final_prompt)
            response_text = gemini_reply.text or "Sorry, I couldn't understand that."
        except Exception as e:
            logging.error(f"Gemini API error: {e}")
            response_text = f"Gemini API error: {e}"

        save_chat_to_db(get_session_id(request), user_msg, response_text)
        return JsonResponse({"response": response_text})
